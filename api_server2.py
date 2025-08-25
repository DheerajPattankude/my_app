import streamlit as st
from docx import Document
from io import BytesIO
import os
from openai import OpenAI
from deep_translator import GoogleTranslator

# =========================
# PAGE CONFIG & STYLE
# =========================
st.set_page_config(page_title="Multi-Agent AI", page_icon="🤖", layout="centered")

# Custom CSS for styling
st.markdown(
    """
    <style>
    body {
        background: linear-gradient(180deg, #f7f9fc, #e0f7fa);
    }
    .stTextArea textarea {
        background-color: #fff8e1 !important;
        border-radius: 10px !important;
        border: 2px solid #ff9800 !important;
        color: #000000 !important;
        font-size: 16px !important;
        padding: 12px !important;
        box-shadow: 0px 3px 8px rgba(0,0,0,0.1);
    }
    .ask-section {
        background: linear-gradient(135deg, #ffe259, #ffa751);
        padding: 20px;
        border-radius: 12px;
        box-shadow: 0px 4px 15px rgba(0,0,0,0.1);
        margin-bottom: 20px;
    }
    .answer-section {
        padding: 15px;
        border-radius: 12px;
        margin-bottom: 20px;
        color: #000;
        box-shadow: 0px 4px 10px rgba(0,0,0,0.05);
    }
    .answer-edu { background: linear-gradient(135deg, #d4fc79, #96e6a1); }
    .answer-police { background: linear-gradient(135deg, #fddb92, #d1fdff); }
    .answer-krishna,
    .answer-Dr_Ambedkar,
    .answer-Bhagwan_Mahaveer,
    .answer-Bhagwan_Budda,
    .answer-IAS_DC,
    .answer-IAS_Secretary {
        background: linear-gradient(135deg, #ffecd2, #fcb69f);
    }
    </style>
    """,
    unsafe_allow_html=True
)

# =========================
# LANGUAGE CODE MAPPING
# =========================
LANGUAGE_CODES = {
    "English": "en",
    "Hindi": "hi",
    "Marathi": "mr",
    "Gujarati": "gu",
    "Tamil": "ta",
    "Telugu": "te",
    "Kannada": "kn",
    "Malayalam": "ml",
    "Bengali": "bn",
    "Punjabi": "pa"
}

# =========================
# SYSTEM PROMPTS FOR AGENTS
# =========================
AGENT_PROMPTS = {
    "Indian Institution Advisor": "you are a best experience lowyer of Indian constitution, give answer in simple language with IPC section and article, also refer previous case study , have references.",
    "Police Guideline Officer": "You are a police guideline officer in India, Give advice based on Indian law, safety, and real-world procedures.",
    "Lord Krishna": "You are Lord Krishna, answering with wisdom from the Bhagavad Gita, Mahabharata, The Vishnu Purana, Bhagavata Purana, Narada Purana, Garuda Purana, and Vayu Purana in a compassionate tone. Answer will be best and short.",
    "Dr. Ambedkar": "You are Dr. Ambedkar, give answers on your life experience and present time condition with also help of constitution. Answer will be best and short.",
    "Bhagwan Mahaveer": "You are Bhagwan Mahaveer, answering with wisdom from the Jain Agamas or Agam Sutras, teachings of Lord Mahavira, Ang-agams, Upang-agams and Darshan shastra like all holy books in a compassionate tone. Answer will be best and short.",
    "Bhagwan Budda": "You are Bhagwan Gautam Buddha, answering with wisdom from The Tripitaka, Vinaya Pitaka, Sutta Pitaka and Abhidhamma Pitaka in a compassionate tone. Answer will be best and short.",
    "IAS role as DC": "You are an IAS officer acting as DC of Indian districts, answer a question as IAS DC of India, you have all power of this role, take decision and give best answer,no need to hide, Answer will be best and short.",
    "IAS role as Secretary": "You are an IAS officer acting as Secretary of Indian states, answer a question as IAS Secretary of India, you have all power of this role, take decision and give best answer, no need to hide, Answer will be best and short."
}

# =========================
# OPENAI CLIENT
# =========================
client = OpenAI(
    base_url="https://router.huggingface.co/v1",
    api_key=os.environ["HF_TOKEN"],
)

def query_model_combined(user_prompt, selected_agents):
    """Make one API call for all selected agents with their prompts."""
    try:
        # Build combined system + user instructions
        combined_instructions = "You are a multi-role assistant. For each role, follow its unique system prompt and answer separately. Format strictly:\n\n### <Role Name>:\nAnswer...\n\n"

        role_texts = []
        for agent in selected_agents:
            sys_prompt = AGENT_PROMPTS[agent]
            role_texts.append(f"Role: {agent}\nSystem Prompt: {sys_prompt}\n")

        final_prompt = (
            f"Question: {user_prompt}\n\n"
            f"Provide answers for the following roles:\n\n" +
            "\n".join(role_texts)
        )

        response = client.chat.completions.create(
            model="meta-llama/Llama-3.1-8B-Instruct:cerebras",
            messages=[
                {"role": "system", "content": combined_instructions},
                {"role": "user", "content": final_prompt}
            ],
            temperature=0.7
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error: {str(e)}"

def translate_text(text, target_lang):
    """Translate text to target language using deep_translator."""
    try:
        return GoogleTranslator(source="auto", target=target_lang).translate(text)
    except Exception as e:
        return f"Translation Error: {str(e)}"

# =========================
# UI INPUT SECTION
# =========================
st.title("🤖 Me CM Assistant")

st.markdown('<div class="ask-section">', unsafe_allow_html=True)

with st.form("query_form"):
    user_question = st.text_area("💬 Ask your question:", height=120)
    
    # Language selection
    language = st.selectbox("🌐 Select language:", list(LANGUAGE_CODES.keys()))
    
    # Agent selection
    agent_options = list(AGENT_PROMPTS.keys())
    selected_agents = st.multiselect(
        "🤝 Select agents to query:", options=agent_options
    )

    submitted = st.form_submit_button("🚀 Get Answers")

st.markdown('</div>', unsafe_allow_html=True)

# =========================
# PROCESSING
# =========================
if submitted and user_question and selected_agents:
    doc = Document()
    doc.add_heading("AI Agent Responses", level=1)
    lang_code = LANGUAGE_CODES[language]

    # ONE API CALL
    raw_response = query_model_combined(user_question, selected_agents)

    # Parse by role (### markers)
    responses = {}
    for block in raw_response.split("### "):
        if not block.strip():
            continue
        try:
            role, answer = block.split(":", 1)
            responses[role.strip()] = answer.strip()
        except:
            continue

    # Display each agent’s response
    for agent in selected_agents:
        answer = responses.get(agent, "No answer generated.")
        translated_answer = translate_text(answer, lang_code)

        css_class = (
            "answer-edu" if agent == "Indian Institution Advisor" else
            "answer-police" if agent == "Police Guideline Officer" else
            "answer-krishna" if agent == "Lord Krishna" else
            "answer-Dr_Ambedkar" if agent == "Dr. Ambedkar" else
            "answer-Bhagwan_Mahaveer" if agent == "Bhagwan Mahaveer" else
            "answer-Bhagwan_Budda" if agent == "Bhagwan Budda" else
            "answer-IAS_DC" if agent == "IAS role as DC" else
            "answer-IAS_Secretary"
        )

        st.markdown(f'<div class="answer-section {css_class}">', unsafe_allow_html=True)
        st.subheader(agent)
        st.write(translated_answer)
        st.markdown('</div>', unsafe_allow_html=True)

        # Add to Word doc
        doc.add_heading(agent, level=2)
        doc.add_paragraph(translated_answer)

    # =========================
    # DOWNLOAD WORD FILE
    # =========================
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button(
        label="📥 Download as Word",
        data=buffer,
        file_name="AI_Agent_Responses.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
